import json
import pandas as pd
import streamlit as st
from datetime import datetime
from typing import Dict, List, Any, Optional
from io import BytesIO
import base64

# Import config for LLM calls
from config_models import call_selected_llm, translate_to_turkish

# For document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

class ReportDataCollector:
    """Analiz verilerini toplayan ve yöneten sınıf"""
    
    def __init__(self):
        self.init_session_state()
    
    def init_session_state(self):
        """Session state'i başlat"""
        if "report_data" not in st.session_state:
            st.session_state.report_data = {
                "dataset_info": {},
                "analysis_history": [],
                "preprocessing_info": "",
                "session_start": datetime.now().isoformat(),
                "total_questions": 0,
                "key_insights": []
            }
    
    def update_dataset_info(self, df: pd.DataFrame, filename: str):
        """Veri seti bilgilerini güncelle"""
        st.session_state.report_data["dataset_info"] = {
            "filename": filename,
            "shape": list(df.shape),
            "columns": df.columns.tolist(),
            "dtypes": df.dtypes.astype(str).to_dict(),
            "memory_usage_mb": round(df.memory_usage(deep=True).sum() / 1024**2, 2),
            "missing_values": df.isnull().sum().to_dict(),
            "numeric_columns": df.select_dtypes(include=['number']).columns.tolist(),
            "categorical_columns": df.select_dtypes(include=['object', 'category']).columns.tolist()
        }
    
    def add_analysis(self, question: str, analysis_type: str, result_summary: str, chart_analysis: str = None):
        """Yeni analiz kaydı ekle"""
        analysis_entry = {
            "timestamp": datetime.now().isoformat(),
            "question": question,
            "analysis_type": analysis_type,  # "visualization", "statistical", "descriptive"
            "result_summary": result_summary,
            "chart_analysis": chart_analysis,
            "question_number": len(st.session_state.report_data["analysis_history"]) + 1
        }
        
        st.session_state.report_data["analysis_history"].append(analysis_entry)
        st.session_state.report_data["total_questions"] = len(st.session_state.report_data["analysis_history"])
    
    def update_preprocessing_info(self, preprocessing_report: str):
        """Veri ön işleme bilgilerini güncelle"""
        st.session_state.report_data["preprocessing_info"] = preprocessing_report
    
    def get_report_data(self) -> Dict:
        """Rapor verilerini al"""
        return st.session_state.report_data.copy()
    
    def clear_report_data(self):
        """Rapor verilerini temizle"""
        if "report_data" in st.session_state:
            del st.session_state.report_data
        self.init_session_state()

class ManagerReportAgent:
    """AI destekli yönetici raporu oluşturan agent"""
    
    def __init__(self):
        self.report_template = self._get_report_template()
    
    def _get_report_template(self) -> str:
        """Rapor şablonunu döndür"""
        return """
        # YÖNETİCİ VERİ ANALİZİ RAPORU
        
        ## 1. YÖNETİCİ ÖZETİ
        [Ana bulgular, önemli trendler ve kritik öneriler]
        
        ## 2. VERİ SETİ GENEL BAKIŞ
        [Veri seti özellikleri, boyut, kalite durumu]
        
        ## 3. YAPILAN ANALİZLER VE BULGULAR
        [Her analiz için detaylı bulgular ve yorumlar]
        
        ## 4. ANAHTAR İÇGÖRÜLER
        [En önemli keşfedilen desenler ve trendler]
        
        ## 5. ÖNERİLER VE SONUÇ
        [Eylem önerileri ve karar destekleri]
        """
    
    def generate_report(self, report_data: Dict) -> str:
        """Rapor verilerinden profesyonel rapor oluştur"""
        
        # Analiz verilerini hazırla
        dataset_summary = self._create_dataset_summary(report_data.get("dataset_info", {}))
        analysis_summary = self._create_analysis_summary(report_data.get("analysis_history", []))
        preprocessing_summary = report_data.get("preprocessing_info", "")
        
        # AI prompt oluştur
        prompt = f"""
Sen bir senior veri analisti olarak, yapılan analizler temelinde profesyonel bir yönetici raporu oluşturacaksın.

VERİ SETİ BİLGİLERİ:
{dataset_summary}

VERİ ÖN İŞLEME:
{preprocessing_summary}

YAPILAN ANALİZLER:
{analysis_summary}

Lütfen aşağıdaki formatta TÜRKÇE profesyonel bir yönetici raporu oluştur:

## YÖNETİCİ ÖZETİ
(En önemli 3-4 bulguyu özetleyin. Sayısal veriler kullanın. Karar alıcılar için kritik noktalar.)

## VERİ SETİ GENEL BAKIŞ  
(Veri setinin özellikleri, boyutu, kalitesi hakkında kısa bilgi)

## DETAYLI ANALİZ BULGULARI
(Her analiz sorusu için ayrı başlık altında detaylı bulgular ve yorumlar)

## ANAHTAR İÇGÖRÜLER
(Keşfedilen en önemli desenler, trendler, korelasyonlar)

## ÖNERİLER VE SONUÇ
(Bulgular temelinde somut eylem önerileri)

KURALLAR:
- Profesyonel ve anlaşılır dil kullan
- Sayısal verilerle destekle
- Her bölüm için yeterli detay ver
- İş dünyası terminolojisi kullan
- Concrete öneriler sun
"""
        
        try:
            messages = [
                {"role": "system", "content": "Sen uzman bir veri analisti ve rapor yazarısın. Profesyonel yönetici raporları hazırlıyorsun."},
                {"role": "user", "content": prompt}
            ]
            
            response = call_selected_llm(messages, max_tokens=2000, temperature=0.3)
            return response
            
        except Exception as e:
            return f"Rapor oluşturulurken hata oluştu: {str(e)}\n\nTemel analiz özeti:\n{analysis_summary}"
    
    def _create_dataset_summary(self, dataset_info: Dict) -> str:
        """Veri seti özetini oluştur"""
        if not dataset_info:
            return "Veri seti bilgisi mevcut değil."
        
        summary = f"""
Dosya Adı: {dataset_info.get('filename', 'Bilinmiyor')}
Boyut: {dataset_info.get('shape', [0, 0])[0]:,} satır × {dataset_info.get('shape', [0, 0])[1]} sütun
Hafıza Kullanımı: {dataset_info.get('memory_usage_mb', 0)} MB
Sayısal Sütunlar: {len(dataset_info.get('numeric_columns', []))} adet
Kategorik Sütunlar: {len(dataset_info.get('categorical_columns', []))} adet
Eksik Değer: {sum(dataset_info.get('missing_values', {}).values())} toplam
"""
        return summary
    
    def _create_analysis_summary(self, analysis_history: List[Dict]) -> str:
        """Analiz geçmişi özetini oluştur"""
        if not analysis_history:
            return "Henüz analiz yapılmamış."
        
        summary = f"Toplam {len(analysis_history)} analiz gerçekleştirildi:\n\n"
        
        for i, analysis in enumerate(analysis_history, 1):
            summary += f"{i}. SORU: {analysis.get('question', '')}\n"
            summary += f"   TİP: {analysis.get('analysis_type', '')}\n"
            summary += f"   ÖZET: {analysis.get('result_summary', '')}\n"
            if analysis.get('chart_analysis'):
                summary += f"   GRAFİK ANALİZİ: {analysis.get('chart_analysis', '')}\n"
            summary += "\n"
        
        return summary

class ReportGenerator:
    """PDF ve Word formatında rapor üreten sınıf"""
    
    def __init__(self):
        self.check_dependencies()
    
    def check_dependencies(self):
        """Gerekli kütüphaneleri kontrol et"""
        if not DOCX_AVAILABLE:
            st.warning("⚠️ python-docx kütüphanesi bulunamadı. Word raporu oluşturulamayacak.")
        
        if not REPORTLAB_AVAILABLE:
            st.warning("⚠️ reportlab kütüphanesi bulunamadı. PDF raporu oluşturulamayacak.")
    
    def generate_word_report(self, report_content: str, filename: str) -> Optional[BytesIO]:
        """Word formatında rapor oluştur"""
        if not DOCX_AVAILABLE:
            st.error("Word raporu oluşturmak için python-docx kütüphanesi gerekli.")
            return None
        
        try:
            doc = Document()
            
            # Başlık
            title = doc.add_heading('YÖNETİCİ VERİ ANALİZİ RAPORU', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Tarih
            date_para = doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_page_break()
            
            # İçerik
            lines = report_content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=2)
                elif line:
                    doc.add_paragraph(line)
            
            # BytesIO'ya kaydet
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            st.error(f"Word raporu oluşturulurken hata: {str(e)}")
            return None
    
    def generate_pdf_report(self, report_content: str, filename: str) -> Optional[BytesIO]:
        """PDF formatında rapor oluştur"""
        if not REPORTLAB_AVAILABLE:
            st.error("PDF raporu oluşturmak için reportlab kütüphanesi gerekli.")
            return None
        
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
            
            # Stil tanımlamaları - UTF-8 destekli
            styles = getSampleStyleSheet()
            
            # Türkçe karakter desteği için font ayarları
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                alignment=1,  # Center
                spaceAfter=30,
                fontName='Helvetica-Bold'  # Türkçe destekli font
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                fontName='Helvetica'
            )
            
            # İçerik listesi
            story = []
            
            # Başlık
            clean_title = self._clean_text_for_pdf("YÖNETİCİ VERİ ANALİZİ RAPORU")
            story.append(Paragraph(clean_title, title_style))
            
            clean_date = self._clean_text_for_pdf(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            story.append(Paragraph(clean_date, normal_style))
            story.append(Spacer(1, 20))
            
            # İçerik - Türkçe karakter problemi için encoding
            lines = report_content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('## '):
                    story.append(Spacer(1, 12))
                    # Türkçe karakterleri encode et
                    clean_text = self._clean_text_for_pdf(line[3:])
                    story.append(Paragraph(clean_text, heading_style))
                elif line.startswith('### '):
                    clean_text = self._clean_text_for_pdf(line[4:])
                    story.append(Paragraph(clean_text, styles['Heading3']))
                elif line:
                    clean_text = self._clean_text_for_pdf(line)
                    story.append(Paragraph(clean_text, normal_style))
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            st.error(f"PDF raporu oluşturulurken hata: {str(e)}")
            return None
    
    def _clean_text_for_pdf(self, text: str) -> str:
        """PDF için Türkçe karakterleri temizle/dönüştür"""
        # Türkçe karakter dönüştürme sözlüğü
        turkish_chars = {
            'ı': 'i', 'İ': 'I', 'ğ': 'g', 'Ğ': 'G',
            'ü': 'u', 'Ü': 'U', 'ş': 's', 'Ş': 'S', 
            'ö': 'o', 'Ö': 'O', 'ç': 'c', 'Ç': 'C'
        }
        
        # Karakterleri dönüştür
        for turkish, english in turkish_chars.items():
            text = text.replace(turkish, english)
        
        # HTML özel karakterlerini temizle
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        return text

def display_report_button():
    """Rapor oluşturma butonunu görüntüle"""
    
    # Rapor verilerini kontrol et
    if "report_data" not in st.session_state:
        return
    
    report_data = st.session_state.report_data
    analysis_count = len(report_data.get("analysis_history", []))
    
    if analysis_count == 0:
        return
    
    # Rapor butonu
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col3:
        if st.button("📊 Yönetici Raporu", help=f"{analysis_count} analiz mevcut", use_container_width=True):
            st.session_state.show_report_modal = True
    
    # Modal dialog
    if st.session_state.get("show_report_modal", False):
        with st.container():
            st.markdown("---")
            st.subheader("📄 Yönetici Raporu Oluştur")
            
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.info(f"🔍 **{analysis_count} analiz** temelinde profesyonel rapor oluşturulacak.")
                
                st.write("**Rapor İçeriği:**")
                st.write("• Yönetici Özeti")
                st.write("• Veri Seti Genel Bakış")
                st.write("• Detaylı Analiz Bulgları")
                st.write("• Anahtar İçgörüler")
                st.write("• Öneriler ve Sonuç")
                
                st.markdown("**Rapor formatını seçin:**")
                
                format_col1, format_col2 = st.columns(2)
                
                with format_col1:
                    if st.button("📄 Word (.docx)", use_container_width=True, type="primary"):
                        generate_and_download_report("word")
                
                with format_col2:
                    if st.button("📑 PDF", use_container_width=True, type="primary"):
                        generate_and_download_report("pdf")
                
                if st.button("❌ İptal", use_container_width=True):
                    st.session_state.show_report_modal = False
                    st.rerun()

def generate_and_download_report(format_type: str):
    """Rapor oluştur ve indir"""
    
    with st.spinner(f"🔄 {format_type.upper()} raporu oluşturuluyor..."):
        try:
            # Veri toplama
            collector = ReportDataCollector()
            report_data = collector.get_report_data()
            
            # AI ile rapor oluşturma
            agent = ManagerReportAgent()
            report_content = agent.generate_report(report_data)
            
            # Dosya adı
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            dataset_name = report_data.get("dataset_info", {}).get("filename", "analiz").replace(".csv", "")
            filename = f"Yonetici_Raporu_{dataset_name}_{timestamp}"
            
            # Format'a göre oluştur
            generator = ReportGenerator()
            
            if format_type == "word":
                buffer = generator.generate_word_report(report_content, filename)
                if buffer:
                    st.download_button(
                        label="📥 Word Raporunu İndir",
                        data=buffer,
                        file_name=f"{filename}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.success("✅ Word raporu başarıyla oluşturuldu!")
            
            elif format_type == "pdf":
                buffer = generator.generate_pdf_report(report_content, filename)
                if buffer:
                    st.download_button(
                        label="📥 PDF Raporunu İndir",
                        data=buffer,
                        file_name=f"{filename}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    st.success("✅ PDF raporu başarıyla oluşturuldu!")
            
            # Modal'ı kapat
            st.session_state.show_report_modal = False
            
        except Exception as e:
            st.error(f"❌ Rapor oluşturulurken hata: {str(e)}")

def integrate_report_tracking():
    """Mevcut app.py ile entegrasyon için yardımcı fonksiyonlar"""
    
    # ReportDataCollector instance'ı oluştur
    if 'report_collector' not in st.session_state:
        st.session_state.report_collector = ReportDataCollector()
    
    return st.session_state.report_collector

# Entegrasyon fonksiyonları
def track_dataset_upload(df: pd.DataFrame, filename: str):
    """Veri seti yüklendiğinde çağır"""
    collector = integrate_report_tracking()
    collector.update_dataset_info(df, filename)

def track_analysis_query(question: str, result_type: str, summary: str, chart_analysis: str = None):
    """Analiz yapıldığında çağır"""
    collector = integrate_report_tracking()
    collector.add_analysis(question, result_type, summary, chart_analysis)

def track_preprocessing(preprocessing_report: str):
    """Veri ön işleme yapıldığında çağır"""
    collector = integrate_report_tracking()
    collector.update_preprocessing_info(preprocessing_report)